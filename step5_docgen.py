import json
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

data = json.load(open('final_table.json'))
entries = data['entries']
summary = data['summary']

# ---------- palette ----------
GREEN_DARK = "1E6B52"    # header
GREEN_MED = "2E8B67"
ROW_A = "FFFFFF"
ROW_B = "E9F5EF"         # light green alternate
BORDER = "BFD8CC"
LINK_BLUE = RGBColor(0x0B, 0x5C, 0x8A)
NAME_COLOR = RGBColor(0x14, 0x3A, 0x2C)
COUNTRY_COLOR = RGBColor(0x5A, 0x6A, 0x63)
PROD_COLOR = RGBColor(0x33, 0x33, 0x33)

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)

def set_cell_borders(cell, color=BORDER, sz=4):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        e = OxmlElement(f'w:{edge}')
        e.set(qn('w:val'), 'single')
        e.set(qn('w:sz'), str(sz))
        e.set(qn('w:space'), '0')
        e.set(qn('w:color'), color)
        borders.append(e)
    tcPr.append(borders)

def set_vertical_align(cell, val='center'):
    tcPr = cell._tc.get_or_add_tcPr()
    va = OxmlElement('w:vAlign')
    va.set(qn('w:val'), val)
    tcPr.append(va)

def repeat_header(row):
    trPr = row._tr.get_or_add_trPr()
    th = OxmlElement('w:tblHeader')
    th.set(qn('w:val'), 'true')
    trPr.append(th)

def set_col_widths(table, widths_cm):
    table.autofit = False
    table.allow_autofit = False
    for row in table.rows:
        for i, w in enumerate(widths_cm):
            row.cells[i].width = Cm(w)

def prevent_row_split(row):
    trPr = row._tr.get_or_add_trPr()
    cant = OxmlElement('w:cantSplit')
    cant.set(qn('w:val'), 'true')
    trPr.append(cant)

# ---------- document ----------
doc = Document()

# landscape + margins
sec = doc.sections[0]
sec.orientation = WD_ORIENT.LANDSCAPE
sec.page_width, sec.page_height = sec.page_height, sec.page_width
sec.left_margin = Cm(1.4)
sec.right_margin = Cm(1.4)
sec.top_margin = Cm(1.3)
sec.bottom_margin = Cm(1.3)

# base font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(9)

# ----- Title block -----
h = doc.add_paragraph()
h.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = h.add_run('Better Cotton Initiative — Members Directory')
run.bold = True
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x1E, 0x6B, 0x52)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub.add_run('Retabulated contact directory  •  Updated 26/06/2026')
r2.italic = True
r2.font.size = Pt(10)
r2.font.color.rgb = RGBColor(0x5A, 0x6A, 0x63)
doc.add_paragraph()

# ----- Table -----
headers = ['No.', 'Name of Company and Country', 'Website', 'Email', 'Products']
col_w = [1.2, 7.2, 5.6, 6.2, 6.4]

table = doc.add_table(rows=1, cols=5)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = 'Table Grid'

hdr = table.rows[0]
repeat_header(hdr)
prevent_row_split(hdr)
for i, htext in enumerate(headers):
    c = hdr.cells[i]
    c.text = ''
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i in (0,) else WD_ALIGN_PARAGRAPH.LEFT
    rr = p.add_run(htext)
    rr.bold = True
    rr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    rr.font.size = Pt(10.5)
    shade(c, GREEN_DARK)
    set_cell_borders(c, color=GREEN_DARK, sz=6)
    set_vertical_align(c, 'center')

def add_multiline_bold(cell, companies):
    """companies: list of 'Name\\nCountry' strings. Name bold, country regular italic."""
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    first = True
    for comp in companies:
        parts = comp.split('\n')
        name = parts[0]
        country = parts[1] if len(parts) > 1 else ''
        if not first:
            p = cell.add_paragraph()
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(2)
        rn = p.add_run(name)
        rn.bold = True
        rn.font.color.rgb = NAME_COLOR
        rn.font.size = Pt(9.5)
        if country:
            rc = p.add_run('\n' + country)
            rc.italic = True
            rc.font.color.rgb = COUNTRY_COLOR
            rc.font.size = Pt(8.5)
        first = False

def add_links(cell, items, color):
    cell.text = ''
    p = cell.paragraphs[0]
    if not items:
        return
    txt = ' ;\n'.join(items)
    rr = p.add_run(txt)
    rr.font.color.rgb = color
    rr.font.size = Pt(8.5)

def add_product(cell, products):
    cell.text = ''
    p = cell.paragraphs[0]
    txt = '; '.join(products) if products else ''
    rr = p.add_run(txt)
    rr.font.color.rgb = PROD_COLOR
    rr.font.size = Pt(8.5)

for idx, e in enumerate(entries, start=1):
    row = table.add_row()
    prevent_row_split(row)
    fill = ROW_A if idx % 2 else ROW_B
    cells = row.cells
    # No.
    cells[0].text = ''
    pnum = cells[0].paragraphs[0]
    pnum.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rn = pnum.add_run(str(idx))
    rn.font.size = Pt(8.5)
    # Company + country
    add_multiline_bold(cells[1], e['companies'])
    # Website
    add_links(cells[2], e['websites'], LINK_BLUE)
    # Email
    add_links(cells[3], e['emails'], LINK_BLUE)
    # Products
    add_product(cells[4], e['products'])
    for ci, c in enumerate(cells):
        shade(c, fill)
        set_cell_borders(c)
        set_vertical_align(c, 'top')

set_col_widths(table, col_w)

# ----- Summary block -----
doc.add_paragraph()
sh = doc.add_paragraph()
shr = sh.add_run('Summary of Entries')
shr.bold = True
shr.font.size = Pt(13)
shr.font.color.rgb = RGBColor(0x1E, 0x6B, 0x52)

stable = doc.add_table(rows=5, cols=2)
stable.style = 'Table Grid'
stable.alignment = WD_TABLE_ALIGNMENT.LEFT
srows = [
    ('Entries with website only (no email address)', summary['only_website_no_email']),
    ('Entries with email address only (no website)', summary['only_email_no_website']),
    ('Entries with neither website nor email address', summary['neither']),
    ('Entries with both website and email address', summary['both']),
    ('Total entries', summary['total_entries']),
]
for i, (label, val) in enumerate(srows):
    lc, vc = stable.rows[i].cells
    lc.text = ''
    lp = lc.paragraphs[0]
    lr = lp.add_run(label)
    lr.font.size = Pt(10)
    vc.text = ''
    vp = vc.paragraphs[0]
    vp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    vr = vp.add_run(str(val))
    vr.bold = True
    vr.font.size = Pt(10)
    is_total = (i == len(srows) - 1)
    fill = GREEN_MED if is_total else (ROW_B if i % 2 == 0 else ROW_A)
    for cc in (lc, vc):
        shade(cc, fill)
        set_cell_borders(cc)
        set_vertical_align(cc, 'center')
        if is_total:
            for pr in cc.paragraphs:
                for rn in pr.runs:
                    rn.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    rn.bold = True
stable.columns[0].width = Cm(11)
stable.columns[1].width = Cm(3)
for r in stable.rows:
    r.cells[0].width = Cm(11)
    r.cells[1].width = Cm(3)

# footer note
doc.add_paragraph()
note = doc.add_paragraph()
nr = note.add_run(
    'Notes: Emails were discovered from members\u2019 live websites where not provided in the source, and every '
    'email domain was DNS-verified (retained only where an MX or A record exists). Multiple companies sharing the '
    'same contact email are grouped in a single row; multiple emails/websites for one company are combined and '
    'separated with \u201c;\u201d. Product descriptions are brief summaries derived from each member\u2019s website / '
    'membership category.'
)
nr.italic = True
nr.font.size = Pt(8)
nr.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

out = 'BCI-Members_Directory_Retabulated.docx'
doc.save(out)
print('saved', out, 'with', len(entries), 'entries')
